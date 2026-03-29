from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from pathlib import Path

base = Path(r'c:\Nilang\eduagent-ai claude(2)')
out_dir = base / 'report_output'
img_dir = out_dir / 'images'
out_dir.mkdir(exist_ok=True)

logo = img_dir / 'edunet_logo.png'
fig1 = img_dir / 'fig_1_1_architecture.png'
fig2 = img_dir / 'fig_1_2_student_flow.png'
fig3 = img_dir / 'fig_2_1_admin_modules.png'
fig4 = img_dir / 'fig_3_1_reminder_flow.png'

def set_normal_style(doc):
    sec = doc.sections[0]
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def p(doc, text='', center=False, bold=False, size=12, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def chapter(doc, title):
    p(doc, title.upper(), bold=True, size=16)

def section(doc, title):
    p(doc, title.upper(), bold=True, size=14)

def subsection(doc, title):
    p(doc, title, bold=True, size=12)

def bullets(doc, items):
    for it in items:
        para = doc.add_paragraph('- ' + it)
        para.paragraph_format.left_indent = Inches(0.2)

def add_caption(doc, text):
    c = p(doc, text, center=True, size=11, italic=True)
    return c

doc = Document()
set_normal_style(doc)

# Cover page
p(doc, 'INDUSTRIAL INTERNSHIP REPORT', center=True, bold=True, size=18)
p(doc, 'EduAgent AI - Student Support and Admin Automation Platform', center=True, bold=True, size=14)
p(doc, 'Submitted by', center=True, bold=True, size=13)
p(doc, 'Nilang Vipulkumar Jotaniya', center=True, bold=True, size=14)
p(doc, 'CVMU Enrollment No.: 12202040701108', center=True, size=13)
p(doc, 'In partial fulfillment for the award of the degree of', center=True)
p(doc, 'BACHELOR OF ENGINEERING / TECHNOLOGY', center=True, bold=True, size=14)
p(doc, 'in Computer Engineering', center=True)
p(doc, 'MBIT', center=True, bold=True)
p(doc, 'The Charutar Vidya Mandal (CVM) University, Vallabh Vidyanagar - 388120', center=True)
p(doc, 'April 2026', center=True, bold=True)
doc.add_page_break()

# Certificate and preface sections
chapter(doc, 'Certificate')
p(doc, 'This is to certify that Mr. Nilang Vipulkumar Jotaniya (Enrollment No. 12202040701108), student of Bachelor of Engineering in Computer Engineering, MBIT, The Charutar Vidya Mandal (CVM) University, has successfully carried out internship/project work titled "EduAgent AI - Student Support and Admin Automation Platform" during the academic year 2025-26.')
p(doc, 'The candidate has completed the internship training period from 22/12/2025 to 10/04/2026 and has submitted this report in partial fulfillment of the degree requirements. The work presented is based on practical implementation, testing, and iterative refinement of a role-based student support system with dedicated student and administrator modules.')
p(doc, 'The project progress, technical approach, module implementation, and documentation have been reviewed as per departmental academic guidelines. To the best of our knowledge, the report is an original work carried out by the student under academic and industry mentorship.')
p(doc, 'Internal Guide: Assistant Professor Mr. Dhruv Dalwadi')
p(doc, 'Head of Department: Dr. Gopi Bhatt')
doc.add_page_break()

chapter(doc, 'Company Certificate')
p(doc, 'Company: Edunet Foundation (SAP)')
if logo.exists():
    doc.add_picture(str(logo), width=Inches(2.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
p(doc, 'This is to certify that Nilang Vipulkumar Jotaniya, a student of MBIT, has successfully completed internship/project work in AI-enabled Student Support System Development from 22/12/2025 to 10/04/2026 (Total duration: 16 weeks) under the guidance of Edunet Foundation (SAP) mentor team.')
p(doc, 'Date: 10/04/2026')
doc.add_page_break()

chapter(doc, 'Declaration')
p(doc, 'I, Nilang Vipulkumar Jotaniya (12202040701108), hereby declare that this report is a bonafide record of work carried out by me and no part has been copied without due reference.')
p(doc, 'Signature: ____________________')
p(doc, 'Date: 10/04/2026')
doc.add_page_break()

chapter(doc, 'Acknowledgement')
p(doc, 'I sincerely express my gratitude to my internal guide, Assistant Professor Mr. Dhruv Dalwadi, for his continuous encouragement, timely technical suggestions, and structured review feedback throughout the internship and project implementation phases. His guidance helped me improve my software design decisions, module planning, and quality of documentation.')
p(doc, 'I am equally thankful to Dr. Gopi Bhatt, Head of Department, Computer Engineering, MBIT, for providing the academic environment and support needed to complete this internship work effectively. Her motivation and departmental support enabled me to align this project with institutional expectations and learning outcomes.')
p(doc, 'I extend special thanks to Edunet Foundation (SAP) and the Code Unnati mentor team for selecting me for the 16-week internship and providing an industry-oriented learning path in emerging technologies. Weekly mentor sessions, milestone-based tasks, and project reviews were highly valuable in shaping the final implementation of EduAgent AI.')
p(doc, 'I also acknowledge my classmates, friends, and family members for their support during development, testing, and report preparation. Their encouragement helped me remain consistent during debugging cycles, UI refinements, and final project consolidation.')
doc.add_page_break()

chapter(doc, 'Abstract')
p(doc, 'EduAgent AI is a role-based academic support platform developed to streamline repetitive student support operations and improve access to institutional information. Traditional support channels in many colleges are fragmented across manual desks, notice boards, and scattered documents, which creates delays and inconsistent communication. This project addresses that gap by combining AI-based query handling with structured administrative control.', italic=True)
p(doc, 'The system is implemented using React + TypeScript frontends, a FastAPI backend, MongoDB-based persistence, and a local LLM runtime through Ollama. It provides two separate interfaces: a student portal for enrollment-based login, profile access, fee visibility, reminders, document center, and chat support; and an admin portal for FAQs, escalations, exams, fee records, student management, and document publishing. The design emphasizes clean role separation, traceable actions, and practical daily usability for institutional environments.', italic=True)
p(doc, 'During development, multiple workflow-level improvements were completed, including student-specific navigation hardening, reminder visibility alignment, document download tracking, and session isolation between admin and student contexts. Testing results show stable operation for core use-cases such as authentication, query handling, reminder delivery, and data updates. The outcome demonstrates a deployable foundation for a college-ready assistant platform with clear future extensibility for analytics, stronger identity management, and multilingual response support.', italic=True)
doc.add_page_break()

chapter(doc, 'List of Figures')
for x in ['Fig 1.1 Overall System Architecture', 'Fig 1.2 Student Portal Navigation Flow', 'Fig 2.1 Admin Dashboard Module Layout', 'Fig 3.1 Reminder Notification Workflow']:
    p(doc, x)

chapter(doc, 'List of Tables')
for x in ['Table 1.1 Technology Stack and Purpose', 'Table 2.1 Functional Module Mapping', 'Table 3.1 API Endpoint Summary', 'Table 4.1 Test Cases and Results', 'Table 4.2 Defects and Fixes']:
    p(doc, x)
doc.add_page_break()

chapter(doc, 'Table of Contents')
for x in [
    'Chapter 1 Introduction to Internship and Project Overview',
    'Chapter 2 Existing System Study and Problem Definition',
    'Chapter 3 Proposed System Architecture and Design',
    'Chapter 4 Implementation and Module Details',
    'Chapter 5 Testing, Validation and Result Analysis',
    'Chapter 6 Conclusion, Limitations and Future Enhancements',
    'References',
    'Appendix'
]:
    p(doc, x)
doc.add_page_break()

chapter(doc, 'Chapter 1 Introduction to Internship and Project Overview')
section(doc, '1.1 Internship Context')
p(doc, 'The internship/project was carried out with the intent to build a production-oriented academic assistant platform for college use. The practical target was to reduce repetitive support dependency and establish a maintainable student-admin workflow under clear role boundaries.')
section(doc, '1.2 Problem Statement')
p(doc, 'Academic support requests related to fees, exams, attendance, and document access are repetitive in nature but often processed manually. This causes delays, inconsistency, and communication overhead. A centralized AI-assisted system with admin governance can resolve this operational gap.')
section(doc, '1.3 Objectives')
bullets(doc, [
    'Enable enrollment-based student authentication and dashboard access.',
    'Provide AI-assisted student chat for institutional questions.',
    'Provide admin control over FAQs, escalations, documents, exams, and fees.',
    'Track reminders and document downloads for visibility and accountability.',
    'Maintain clean role boundaries between student and admin capabilities.'
])
section(doc, '1.4 Technology Stack')
t1 = doc.add_table(rows=1, cols=3)
for i,h in enumerate(['Layer', 'Technology', 'Purpose']):
    t1.rows[0].cells[i].text = h
for row in [
    ('Frontend', 'React + TypeScript + Vite', 'Student and Admin web interfaces'),
    ('Backend', 'FastAPI (Python)', 'Role-aware APIs and orchestration'),
    ('Database', 'MongoDB', 'Persistent data and event storage'),
    ('LLM Runtime', 'Ollama (phi3:mini)', 'Local response generation')
]:
    c = t1.add_row().cells
    c[0].text, c[1].text, c[2].text = row
if fig1.exists():
    add_caption(doc, 'Fig 1.1 Overall System Architecture')
    doc.add_picture(str(fig1), width=Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
if fig2.exists():
    add_caption(doc, 'Fig 1.2 Student Portal Navigation Flow')
    doc.add_picture(str(fig2), width=Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

chapter(doc, 'Chapter 2 Existing System Study and Problem Definition')
section(doc, '2.1 Existing Process Analysis')
p(doc, 'The current environment relies on fragmented sources such as notice boards, office interactions, PDF circulars, and informal communication channels. Students often do not know where to find reliable and updated information.')
section(doc, '2.2 Identified Weaknesses')
bullets(doc, [
    'Repeated queries consume substantial administrative effort.',
    'Response quality varies by communication channel.',
    'No standard escalation lifecycle for unresolved cases.',
    'No measurable tracking for document access and reminder reach.'
])
section(doc, '2.3 Requirement Mapping')
subsection(doc, 'Functional Requirements')
bullets(doc, [
    'Student login and profile visibility',
    'Student chat and response feedback',
    'Admin FAQ/escalation/document/exam/fee/student management',
    'Reminder notifications and read status tracking'
])
subsection(doc, 'Non-Functional Requirements')
bullets(doc, [
    'Role safety and predictable behavior',
    'Simple UI for frequent daily use',
    'Maintainable codebase and modular backend APIs'
])
if fig3.exists():
    add_caption(doc, 'Fig 2.1 Admin Dashboard Module Layout')
    doc.add_picture(str(fig3), width=Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

chapter(doc, 'Chapter 3 Proposed System Architecture and Design')
section(doc, '3.1 Architecture Overview')
p(doc, 'The proposed architecture uses two frontend applications (student and admin) communicating with a shared FastAPI backend. The backend coordinates MongoDB for persistent records and Ollama for generated responses. Design emphasizes separation of concerns and controlled role access.')
section(doc, '3.2 Functional Modules')
t2 = doc.add_table(rows=1, cols=3)
for i,h in enumerate(['Module', 'Actor', 'Description']):
    t2.rows[0].cells[i].text = h
for row in [
    ('Student Login', 'Student', 'Enrollment/password authentication'),
    ('Student Details', 'Student', 'Profile, fees, notices, documents'),
    ('Student Chat', 'Student', 'Academic assistant conversation'),
    ('FAQ Management', 'Admin', 'Knowledge base CRUD and updates'),
    ('Escalation Handling', 'Admin', 'Status tracking and resolution flow'),
    ('Document Library', 'Admin', 'Upload, list, and monitor downloads'),
    ('Student Management', 'Admin', 'Create, edit, import student records'),
    ('Fee Ledger', 'Admin', 'Due/balance tracking and reminders')
]:
    c = t2.add_row().cells
    c[0].text, c[1].text, c[2].text = row
section(doc, '3.3 Security and Access Design')
bullets(doc, [
    'Student session token governs student-only APIs.',
    'Admin auth governs mutation operations in admin modules.',
    'Student UI excludes admin authority controls after login.',
    'Fresh-open query parameter resets student token when opened from admin.'
])
if fig4.exists():
    add_caption(doc, 'Fig 3.1 Reminder Notification Workflow')
    doc.add_picture(str(fig4), width=Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

chapter(doc, 'Chapter 4 Implementation and Module Details')
section(doc, '4.1 Backend Implementation')
p(doc, 'The backend is implemented with FastAPI and organized around role-specific endpoint groups. Student endpoints handle authentication, profile retrieval, fee data, reminders, and document access. Admin endpoints manage operational data entities and state transitions.')
section(doc, '4.2 API Endpoint Summary')
t3 = doc.add_table(rows=1, cols=4)
for i,h in enumerate(['Endpoint', 'Method', 'Role', 'Purpose']):
    t3.rows[0].cells[i].text = h
for row in [
    ('/api/student/login', 'POST', 'Student', 'Authenticate student'),
    ('/api/student/me', 'GET', 'Student', 'Fetch profile data'),
    ('/api/student/fees', 'GET', 'Student', 'Fetch ledger rows'),
    ('/api/student/reminders', 'GET', 'Student', 'Fetch notices'),
    ('/api/student/documents', 'GET', 'Student', 'Document list'),
    ('/api/chat', 'POST', 'Student', 'AI response generation'),
    ('/api/admin/faqs', 'GET/POST', 'Admin', 'FAQ management'),
    ('/api/admin/students', 'GET/POST', 'Admin', 'Student management'),
    ('/api/admin/students/{id}', 'PUT', 'Admin', 'Edit student record'),
    ('/api/admin/pdfs', 'POST', 'Admin', 'Upload document'),
    ('/api/admin/fees/reminder', 'POST', 'Admin', 'Send reminder')
]:
    c = t3.add_row().cells
    c[0].text, c[1].text, c[2].text, c[3].text = row
section(doc, '4.3 Frontend Behavior and Routing')
bullets(doc, [
    'Student app routes: /details and /chat after login.',
    'Default post-login route targets student details.',
    'Student sidebar now includes only student sections and back-to-admin navigation link.',
    'Admin panel and student panel are isolated by role.'
])
section(doc, '4.4 Final UI/Workflow Fixes')
bullets(doc, [
    'Removed admin panel access from student operations sidebar.',
    'Added back-to-admin link in student panel for navigation convenience only.',
    'Implemented fresh student login opening from admin launcher.',
    'Aligned login typography and corrected helper-text spacing.'
])

doc.add_page_break()

chapter(doc, 'Chapter 5 Testing, Validation and Result Analysis')
section(doc, '5.1 Test Strategy')
bullets(doc, [
    'Functional flow validation for both portals',
    'Role access and routing validation',
    'Build and compile checks',
    'Regression checks for recent bug fixes'
])
section(doc, '5.2 Test Cases and Results')
t4 = doc.add_table(rows=1, cols=5)
for i,h in enumerate(['Test ID', 'Condition', 'Expected Output', 'Actual Output', 'Status']):
    t4.rows[0].cells[i].text = h
for row in [
    ('TC-01', 'Valid student login', 'Student dashboard opens', 'Matched', 'Pass'),
    ('TC-02', 'Student sidebar options', 'No admin control item', 'Matched', 'Pass'),
    ('TC-03', 'Open student from admin', 'Fresh login page opens', 'Matched', 'Pass'),
    ('TC-04', 'Admin edits student details', 'Record updates persisted', 'Matched', 'Pass'),
    ('TC-05', 'Admin reminder send', 'Student reminder visible', 'Matched', 'Pass'),
    ('TC-06', 'Document upload/download', 'Download events recorded', 'Matched', 'Pass'),
    ('TC-07', 'FAQ helpful click', 'Counters update correctly', 'Matched', 'Pass')
]:
    c = t4.add_row().cells
    for i,v in enumerate(row):
        c[i].text = v
section(doc, '5.3 Defect and Fix Log')
t5 = doc.add_table(rows=1, cols=3)
for i,h in enumerate(['Issue', 'Cause', 'Fix']):
    t5.rows[0].cells[i].text = h
for row in [
    ('Admin option seen in student panel', 'Navigation regression', 'Restored student-only sidebar entries'),
    ('Student session auto-carried to new tab', 'Persisted token reuse', 'Fresh-open flow clears token before init'),
    ('Login helper line looked merged', 'Inline text rendering mismatch', 'Adjusted helper block layout and spacing')
]:
    c = t5.add_row().cells
    c[0].text, c[1].text, c[2].text = row
section(doc, '5.4 Validation Summary')
bullets(doc, [
    'Student UI build: Pass',
    'Admin UI build: Pass',
    'Backend Python compile checks: Pass',
    'Critical role boundary and navigation issues resolved'
])

doc.add_page_break()

chapter(doc, 'Chapter 6 Conclusion, Limitations and Future Enhancements')
section(doc, '6.1 Conclusion')
p(doc, 'EduAgent AI successfully delivers a practical student support and admin operations platform for college workflows. The final implementation achieves role-safe navigation, manageable data operations, and student-facing assistance with measurable administrative visibility.')
section(doc, '6.2 Limitations')
bullets(doc, [
    'LLM response speed depends on local machine resources',
    'Production-grade identity federation not implemented yet',
    'Advanced institutional analytics not yet included'
])
section(doc, '6.3 Future Enhancements')
bullets(doc, [
    'Single sign-on and stronger security policy controls',
    'Streaming response UX for chat',
    'Multi-language support',
    'Admin analytics dashboard for decision support',
    'Automated semester-wise reporting exports'
])

doc.add_page_break()

chapter(doc, 'References')
for r in [
    'FastAPI Documentation - https://fastapi.tiangolo.com',
    'React Documentation - https://react.dev',
    'MongoDB Documentation - https://www.mongodb.com/docs',
    'Vite Documentation - https://vitejs.dev',
    'Ollama Documentation - https://ollama.com',
    'TypeScript Documentation - https://www.typescriptlang.org/docs/'
]:
    p(doc, r)

doc.add_page_break()
chapter(doc, 'Appendix')
section(doc, 'Appendix A - Screenshot Index')
bullets(doc, [
    'Student Login Screen',
    'Student Details Screen',
    'Student Chat Screen',
    'Admin Dashboard Overview',
    'Admin FAQ Management',
    'Admin Student Management'
])
section(doc, 'Appendix B - Internship Experience Summary')
p(doc, 'As per the Edunet Foundation internship offer dated 09/02/2026, I was selected for the Code Unnati program internship track focused on emerging technologies, including advanced competitive coding, Artificial Intelligence, Machine Learning, and SAP technical skill orientation. The internship duration was 16 weeks, beginning from 22/12/2025, and required regular participation in mentor-led weekly sessions.')
p(doc, 'The internship structure was milestone-driven. I attended scheduled mentoring sessions, completed weekly tasks, and aligned project work with the prescribed learning plan. Mandatory criteria included consistent mentor interaction, timely milestone completion, and final project presentation. This structure helped maintain a disciplined engineering workflow and improved both technical depth and communication clarity.')
p(doc, 'Within this framework, my primary project contribution was the end-to-end development of EduAgent AI, a role-separated student and admin support platform. My responsibilities included requirement analysis, backend API design, database modeling, role-aware access control, frontend workflow integration, defect resolution, and iterative UI refinements for practical academic usage.')
p(doc, 'A key learning outcome from the internship was translating institution-level support problems into maintainable software modules. I improved in designing data models for students, reminders, and documents; implementing controlled CRUD operations; integrating AI response pipelines; and validating role boundaries to prevent unauthorized feature exposure. The debugging and optimization phase significantly strengthened my problem-solving approach.')
p(doc, 'The internship also emphasized professional discipline through the undertaking conditions, including attendance, continuous progress visibility, ethical conduct, and accountable communication with mentors. Overall, the experience provided a strong bridge between classroom learning and real implementation practices, and it directly shaped the quality and completeness of this project report.')

out = out_dir / 'Industrial_Internship_Report_Nilang_12202040701108_DETAILED_V2.docx'
doc.save(out)
print(out)
